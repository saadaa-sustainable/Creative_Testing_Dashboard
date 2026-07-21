"""Convert BACKEND_ARCHITECTURE.md to a properly-styled .docx.

Handles:
  - # / ## / ### / #### headings
  - fenced code blocks (```lang ... ```)
  - bulleted lists (- or *) and nested lists (2-space indent)
  - numbered lists (1. 2. …)
  - inline `code`, **bold**, *italic*, [link](url)
  - simple pipe tables (| col1 | col2 |)
  - horizontal rules (---)
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import sys
SRC = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("D:/Creative_Testing_Dashboard/BACKEND_ARCHITECTURE.md")
OUT = SRC.with_suffix('.docx')

# ── styling helpers ─────────────────────────────────────────
CODE_FONT = "Consolas"
CODE_FONT_SIZE = 9
CODE_COLOR = RGBColor(0x22, 0x22, 0x22)
CODE_SHADE = "F5F5F5"

def _shade(cell_or_para, color_hex):
    """Apply background shading to a paragraph or cell."""
    tc_pr = None
    if hasattr(cell_or_para, '_tc'):
        tc_pr = cell_or_para._tc.get_or_add_tcPr()
    else:
        tc_pr = cell_or_para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)

def add_code_block(doc, lang, code_text):
    """Add a single-cell shaded paragraph for a code block."""
    lines = code_text.rstrip('\n').split('\n')
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.right_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(6)
    _shade(p, CODE_SHADE)
    for i, line in enumerate(lines):
        run = p.add_run(line + ('\n' if i < len(lines) - 1 else ''))
        run.font.name = CODE_FONT
        run.font.size = Pt(CODE_FONT_SIZE)
        run.font.color.rgb = CODE_COLOR

# Inline pattern for **bold**, *italic*, `code`, [text](url)
INLINE_RE = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))')

def add_inline_runs(paragraph, text):
    """Render inline formatting: bold, italic, code, links."""
    if not text:
        return
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            r = paragraph.add_run(part[2:-2]); r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = paragraph.add_run(part[1:-1])
            r.font.name = CODE_FONT; r.font.size = Pt(CODE_FONT_SIZE)
            r.font.color.rgb = CODE_COLOR
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            r = paragraph.add_run(part[1:-1]); r.italic = True
        elif part.startswith('['):
            m = re.match(r'\[([^\]]+)\]\(([^)]+)\)', part)
            if m:
                r = paragraph.add_run(m.group(1))
                r.font.color.rgb = RGBColor(0x1A, 0x73, 0xE8)
                r.underline = True
            else:
                paragraph.add_run(part)
        else:
            paragraph.add_run(part)

def add_heading(doc, level, text):
    hdr = doc.add_heading('', level=level)
    add_inline_runs(hdr, text)

def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single'); bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1'); bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom); pPr.append(pBdr)

def add_table_from_pipes(doc, header_line, rows_lines):
    """Render a markdown pipe table."""
    def cells(line):
        return [c.strip() for c in line.strip().strip('|').split('|')]
    hdr = cells(header_line)
    body = [cells(r) for r in rows_lines]
    tbl = doc.add_table(rows=1 + len(body), cols=len(hdr))
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    for i, h in enumerate(hdr):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        add_inline_runs(p, h)
        for run in p.runs: run.bold = True
    for r_i, row in enumerate(body, start=1):
        for c_i, val in enumerate(row):
            cell = tbl.rows[r_i].cells[c_i]
            cell.text = ''
            add_inline_runs(cell.paragraphs[0], val)

# ── main parse ─────────────────────────────────────────
doc = Document()

# Sensible default font sizing (headings styled by python-docx)
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

with open(SRC, encoding='utf-8') as f:
    lines = f.read().split('\n')

i = 0
in_code = False
code_lang = ''
code_buf = []

while i < len(lines):
    line = lines[i]

    # Fenced code blocks
    if line.startswith('```'):
        if in_code:
            add_code_block(doc, code_lang, '\n'.join(code_buf))
            code_buf = []
            in_code = False
        else:
            code_lang = line[3:].strip()
            in_code = True
        i += 1
        continue

    if in_code:
        code_buf.append(line)
        i += 1
        continue

    # Horizontal rule
    if line.strip() == '---':
        add_hr(doc)
        i += 1
        continue

    # Headings
    m = re.match(r'^(#{1,4})\s+(.*)$', line)
    if m:
        level = len(m.group(1))
        add_heading(doc, level, m.group(2).strip())
        i += 1
        continue

    # Pipe table
    if line.strip().startswith('|') and i + 1 < len(lines) and re.match(r'^\|\s*:?-+', lines[i+1].strip()):
        header = line
        j = i + 2
        rows = []
        while j < len(lines) and lines[j].strip().startswith('|'):
            rows.append(lines[j]); j += 1
        add_table_from_pipes(doc, header, rows)
        i = j
        continue

    # List items (numeric and bullet)
    m_bullet = re.match(r'^(\s*)([-*])\s+(.*)$', line)
    m_number = re.match(r'^(\s*)(\d+)\.\s+(.*)$', line)
    if m_bullet or m_number:
        m = m_bullet or m_number
        indent = len(m.group(1))
        text = m.group(3)
        # Merge continuation lines (indented, not another list item)
        while (i + 1 < len(lines)
               and lines[i+1].strip()
               and not re.match(r'^\s*[-*\d]', lines[i+1])
               and not lines[i+1].startswith('#')
               and not lines[i+1].startswith('```')
               and not lines[i+1].strip().startswith('|')):
            text += ' ' + lines[i+1].strip()
            i += 1
        style_name = 'List Number' if m_number else 'List Bullet'
        try:
            p = doc.add_paragraph(style=style_name)
        except KeyError:
            p = doc.add_paragraph()
            p.style = doc.styles['Normal']
        if indent:
            p.paragraph_format.left_indent = Cm(0.8 + 0.6 * (indent // 2))
        add_inline_runs(p, text)
        i += 1
        continue

    # Blank line
    if not line.strip():
        i += 1
        continue

    # Plain paragraph — merge subsequent non-blank non-special lines
    para = line
    while (i + 1 < len(lines)
           and lines[i+1].strip()
           and not lines[i+1].startswith('#')
           and not lines[i+1].startswith('```')
           and not lines[i+1].strip().startswith('|')
           and not re.match(r'^\s*[-*\d]+[.\s]', lines[i+1])
           and lines[i+1].strip() != '---'):
        para += ' ' + lines[i+1].strip()
        i += 1
    p = doc.add_paragraph()
    add_inline_runs(p, para)
    i += 1

# Title page tweak — bump first heading up to Title style if it's an h1
if doc.paragraphs and doc.paragraphs[0].style.name.startswith('Heading 1'):
    doc.paragraphs[0].style = doc.styles['Title']

doc.save(OUT)
print(f"wrote {OUT}")
