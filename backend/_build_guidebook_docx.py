"""
_build_guidebook_docx.py — render the dashboard, take screenshots, build a .docx guidebook.

Loads Supabase creds from .env via dotenv (never printed). Derives the REST
URL from SUPABASE_DB_URL and uses the service_role key for API auth (RLS is
already permissive on every dashboard table, so anon vs service_role look
identical). Injects the two creds via URL params to the LOCAL index_v2.html
file, then walks each view and screenshots key states.

Output:
  D:/Creative_Testing_Dashboard/GUIDEBOOK.docx
"""
from __future__ import annotations
import os, re, sys, time, urllib.parse
from pathlib import Path
from dotenv import load_dotenv
from playwright.sync_api import sync_playwright, Page
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ────────────────────────────────────────────────────────────────────
# Credential loading — anything with these substrings gets scrubbed
# before ANY print. Screenshots are viewport-only (no URL bar).
# ────────────────────────────────────────────────────────────────────
try: sys.stdout.reconfigure(encoding="utf-8", errors="backslashreplace")
except Exception: pass
load_dotenv(override=True)

DB_URL = (os.environ.get("SUPABASE_DB_URL") or "").strip()
KEY    = (os.environ.get("service_role") or os.environ.get("SUPABASE_SERVICE_ROLE_KEY") or "").strip()
if not DB_URL or not KEY:
    raise SystemExit("Missing SUPABASE_DB_URL or service_role in .env")

# Extract project ref from postgres URL. Handles both:
#   postgresql://postgres:PW@db.<REF>.supabase.co:5432/postgres          (direct)
#   postgresql://postgres.<REF>:PW@aws-...-pooler.supabase.com:5432/...  (pooler)
def _extract_ref(url: str) -> str:
    for pat in (r"db\.([a-z0-9]+)\.supabase\.co",
                r"postgres\.([a-z0-9]+):",
                r"@([a-z0-9]+)\.supabase\.co"):
        m = re.search(pat, url)
        if m: return m.group(1)
    return ""
PROJECT_REF = _extract_ref(DB_URL)
if not PROJECT_REF:
    raise SystemExit("Could not extract project ref from SUPABASE_DB_URL "
                     "(neither db.<ref>.supabase.co nor postgres.<ref>: form matched)")
REST_URL    = f"https://{PROJECT_REF}.supabase.co"

TOK_RE = re.compile(r"(?:eyJ[\w\-.]{40,}|EAA[A-Za-z0-9]{30,}|IGQ[\w\-]{20,})")
def scrub(s: str) -> str:
    s = TOK_RE.sub("<REDACTED-KEY>", str(s or ""))
    s = re.sub(r"(apikey|supabaseAnon|access_token|token)=[^&\s\"'\\]+",
               r"\\1=<REDACTED>", s, flags=re.I)
    return s

def say(*parts):
    print(scrub(" ".join(str(p) for p in parts)))

say(f"[✓] project_ref = {PROJECT_REF}")
say(f"[✓] rest_url    = {REST_URL}")
say(f"[✓] key         = <REDACTED, {len(KEY)} chars>")

# ────────────────────────────────────────────────────────────────────
# Paths
# ────────────────────────────────────────────────────────────────────
ROOT   = Path(__file__).parent.parent
HTML   = ROOT / "index_v2.html"
SHOTS  = ROOT / "docs" / "_guidebook_shots"
GUIDE  = ROOT / "GUIDEBOOK.md"
DOCX   = ROOT / "GUIDEBOOK.docx"

if not HTML.exists():
    raise SystemExit(f"Missing {HTML}")
if not GUIDE.exists():
    raise SystemExit(f"Missing {GUIDE}")
SHOTS.mkdir(parents=True, exist_ok=True)

# ────────────────────────────────────────────────────────────────────
# URL builder — creds live in query params on the LOCAL file URL only
# ────────────────────────────────────────────────────────────────────
def dashboard_url() -> str:
    q = urllib.parse.urlencode({"supabaseUrl": REST_URL, "supabaseAnon": KEY})
    return HTML.resolve().as_uri() + f"?{q}"

# ────────────────────────────────────────────────────────────────────
# Screenshot plan — each entry: (filename_stem, hover_selector_or_None,
#   click_selector_or_None, wait_after_ms, extra_action)
# extra_action is a callable(page) executed AFTER click.
# ────────────────────────────────────────────────────────────────────

def click_if(page: Page, sel: str, timeout: int = 3000):
    try:
        page.wait_for_selector(sel, state="visible", timeout=timeout)
        page.click(sel)
        return True
    except Exception:
        return False

def sc(page: Page, name: str, full_page: bool = False, clip=None):
    p = SHOTS / f"{name}.png"
    page.screenshot(path=str(p), full_page=full_page, clip=clip, animations="disabled")
    say(f"    -> {name}.png ({p.stat().st_size // 1024} kB)")

def build_screenshots():
    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        # Force IST so the dashboard's `new Date()` matches the pipeline's
        # "today" and applyDatePreset('last30') hits the results_table cache
        # window (otherwise we're off by one day and the slow-path kicks in).
        ctx = browser.new_context(viewport={"width": 1600, "height": 950},
                                  device_scale_factor=1.5,
                                  timezone_id="Asia/Kolkata",
                                  locale="en-IN")
        page = ctx.new_page()
        say(f"[*] opening dashboard  (viewport 1600x950 @1.5x)")
        # log console errors so silent fetch failures surface
        page.on("console", lambda m: (
            say(f"    [console.{m.type}]", scrub(m.text)[:200])
            if m.type in ("error","warning") else None
        ))
        try:
            page.goto(dashboard_url(), wait_until="domcontentloaded",
                      timeout=30_000)
        except Exception as e:
            say(f"    ! goto failed: {scrub(str(e))[:180]}")
            browser.close(); return
        # Playwright's headless browser computes "today" as 2026-07-01
        # (UTC boundary) even with Asia/Kolkata locale, so the default
        # Last-30d window falls a day off the results_table cache. Give
        # the init IIFE a moment to run, then click a preset that maps
        # to a smaller live-aggregation window (2 days) so it finishes
        # quickly regardless of cache alignment.
        page.wait_for_timeout(2500)
        page.evaluate("""
          () => {
            const btn = document.querySelector('.preset[data-p="thisMonth"]');
            if (btn) btn.click();
          }
        """)
        # Wait for data by observing DOM signals — allAds and primaryAds
        # are `let`-scoped so they're not on window and page.evaluate
        # can't see them directly. Instead we watch the KPI tile #kp-winner
        # for a non-"—" text (means renderKpis has run against non-empty
        # primaryAds) and the dbStat strip for a "Loaded" prefix (means
        # results_table cache / primary aggregation completed).
        say("    waiting for data …")
        try:
            page.wait_for_function(
                """() => {
                    const kp = document.querySelector('#kp-winner');
                    const db = document.querySelector('#dbStat');
                    const kpReady = kp && kp.textContent && kp.textContent.trim() !== '—';
                    const dbReady = db && /Loaded/i.test(db.textContent || '');
                    return kpReady && dbReady;
                }""",
                timeout=180_000
            )
            page.wait_for_timeout(3000)   # let charts + funnel render
            db_text = page.evaluate("document.querySelector('#dbStat')?.textContent || ''")
            say(f"    ready — {scrub(db_text).strip()}")
        except Exception as e:
            db_text = page.evaluate("document.querySelector('#dbStat')?.textContent || ''")
            say(f"    ! data-load wait timed out — dbStat: {scrub(db_text).strip()}")

        # ── 1. Creative Testing (default view)
        say("[*] Creative Testing")
        page.evaluate("window.scrollTo(0,0)")
        page.wait_for_timeout(400)
        sc(page, "01_ct_overview")
        # zoom in on the toolbar + KPI strip
        sc(page, "01a_ct_kpis", clip={"x":220,"y":80,"width":1370,"height":320})
        # scroll to the funnel + focus panels
        page.evaluate("document.querySelector('#funnelBody')?.scrollIntoView({block:'center'})")
        page.wait_for_timeout(400)
        sc(page, "01b_ct_funnel")

        # ── 2. Creative Lifecycle
        say("[*] Creative Lifecycle")
        page.click('.sb-item[data-view="lifecycle"]')
        page.wait_for_timeout(1500)
        page.evaluate("window.scrollTo(0,0)")
        sc(page, "02_lifecycle_overview")
        # click the 3× bucket to open the drill
        clicked = page.evaluate("""
          () => {
            const c = document.querySelector('.freq-bucket[data-b=\"b5\"]');
            if (!c) return false; c.click(); return true;
          }
        """)
        if clicked:
            page.wait_for_timeout(1200)
            page.evaluate("document.getElementById('freqDrill')?.scrollIntoView({block:'center'})")
            page.wait_for_timeout(400)
            sc(page, "02a_lifecycle_freq_drill")

        # ── 3. Ads Analyse
        say("[*] Ads Analyse")
        page.click('.sb-item[data-view="ae"]')
        page.wait_for_timeout(2000)
        page.evaluate("window.scrollTo(0,0)")
        sc(page, "03_ae_overview")
        # scroll a bit to show the main table + filters together
        page.evaluate("window.scrollBy(0, 300)")
        page.wait_for_timeout(300)
        sc(page, "03a_ae_table")

        # ── 4. Ad Intelligence
        say("[*] Ad Intelligence")
        page.click('.sb-item[data-view="adintel"]')
        page.wait_for_timeout(3000)  # this view fetches on nav
        page.evaluate("window.scrollTo(0,0)")
        sc(page, "04_adintel_overview")

        # ── 5. Inventory
        say("[*] Inventory")
        page.click('.sb-item[data-view="inventory"]')
        page.wait_for_timeout(2500)
        page.evaluate("window.scrollTo(0,0)")
        sc(page, "05_inventory_overview")

        # ── 6. Definitions modal (back to Creative Testing)
        say("[*] Definitions modal")
        page.click('.sb-item[data-view="testing"]')
        page.wait_for_timeout(1200)
        if click_if(page, '#btnDef'):
            page.wait_for_timeout(800)
            sc(page, "06_definitions_modal")
            # close via escape
            page.keyboard.press("Escape")
            page.wait_for_timeout(400)

        # ── 7. Bucket drill modal (click first Product Focus pill)
        say("[*] Bucket drill modal")
        page.wait_for_timeout(800)
        # Scroll the Product Focus strip into view first — pills only
        # register clicks when visible
        page.evaluate("document.getElementById('prodStrip')?.scrollIntoView({block:'center'})")
        page.wait_for_timeout(400)
        opened = page.evaluate("""
          () => {
            // Prefer the highest-spend product pill (usually 'Collection page')
            const pills = [...document.querySelectorAll('#prodStrip .focus-pill')]
              .filter(p => !p.classList.contains('disabled'));
            if (!pills.length) return {ok:false, reason:'no-pills'};
            pills[0].click();
            return {ok:true, count: pills.length};
          }
        """)
        if opened.get("ok"):
            page.wait_for_selector('#bucketModal', state="visible", timeout=4000)
            page.wait_for_timeout(1800)   # charts finish rendering
            sc(page, "07_bucket_drill_modal")
            page.keyboard.press("Escape")
            page.wait_for_timeout(400)
        else:
            say(f"    ! bucket drill skipped: {opened.get('reason','?')}")

        browser.close()

# ────────────────────────────────────────────────────────────────────
# Markdown → DOCX conversion (subset: headings, code blocks, bullets,
# bold/italic, image tags, plain paras). Not a general MD parser —
# tuned for GUIDEBOOK.md's specific style.
# ────────────────────────────────────────────────────────────────────

# Section -> screenshot mapping (image is inserted at the start of the
# section body). Match by heading TEXT (case-insensitive substring).
SECTION_IMAGES = {
    "getting started":            [],
    "sidebar navigation":         [],
    "section 1 — creative testing":
        ["01_ct_overview.png", "01a_ct_kpis.png", "01b_ct_funnel.png"],
    "section 2 — creative lifecycle":
        ["02_lifecycle_overview.png", "02a_lifecycle_freq_drill.png"],
    "section 3 — ads analyse":
        ["03_ae_overview.png", "03a_ae_table.png"],
    "section 4 — ad intelligence":
        ["04_adintel_overview.png"],
    "section 5 — inventory":
        ["05_inventory_overview.png"],
    "modals and drawers":
        ["06_definitions_modal.png", "07_bucket_drill_modal.png"],
}

INLINE_BOLD   = re.compile(r"\*\*([^*]+)\*\*")
INLINE_ITALIC = re.compile(r"(?<!\*)\*(?!\s)([^*]+?)(?<!\s)\*(?!\*)")
INLINE_CODE   = re.compile(r"`([^`]+)`")
LINK_RE       = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
IMG_RE        = re.compile(r"^!\[[^\]]*\]\(([^)]+)\)\s*$")

def set_font(run, name="Segoe UI", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    r = run._element
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr"); r.insert(0, rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    for a in ("w:ascii","w:hAnsi","w:cs","w:eastAsia"):
        rFonts.set(qn(a), name)
    if size   is not None: run.font.size = Pt(size)
    if bold   is not None: run.font.bold = bold
    if italic is not None: run.font.italic = italic
    if color  is not None: run.font.color.rgb = RGBColor(*color)

def add_shaded_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # grey shading behind the paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F1E6")   # warm sand — matches dashboard chrome
    pPr.append(shd)
    for line in text.split("\n"):
        run = p.add_run(line + "\n")
        set_font(run, "Consolas", size=9.5)
    return p

def add_para_with_inline(doc, text, style=None, size=10.5):
    p = doc.add_paragraph(style=style)
    if not text.strip():
        return p
    # Split by inline formatting tokens; keep order
    tokens = []
    i = 0
    while i < len(text):
        # bold
        mb = INLINE_BOLD.match(text, i)
        mc = INLINE_CODE.match(text, i)
        mi = INLINE_ITALIC.match(text, i)
        ml = LINK_RE.match(text, i)
        if mb:
            tokens.append(("b", mb.group(1))); i = mb.end()
        elif mc:
            tokens.append(("c", mc.group(1))); i = mc.end()
        elif mi:
            tokens.append(("i", mi.group(1))); i = mi.end()
        elif ml:
            tokens.append(("l", ml.group(1))); i = ml.end()
        else:
            # plain text up to next special
            j = i + 1
            while j < len(text):
                if (INLINE_BOLD.match(text, j) or INLINE_CODE.match(text, j)
                    or INLINE_ITALIC.match(text, j) or LINK_RE.match(text, j)):
                    break
                j += 1
            tokens.append(("t", text[i:j])); i = j
    for kind, val in tokens:
        run = p.add_run(val)
        if kind == "b":
            set_font(run, "Segoe UI", size=size, bold=True)
        elif kind == "i":
            set_font(run, "Segoe UI", size=size, italic=True)
        elif kind == "c":
            set_font(run, "Consolas",  size=size - 0.5,
                     color=(0x8A, 0x5A, 0x0A))
        elif kind == "l":
            set_font(run, "Segoe UI", size=size, color=(0x2E, 0x5A, 0xA8))
        else:
            set_font(run, "Segoe UI", size=size)
    return p

def add_screenshot(doc, name):
    p = SHOTS / name
    if not p.exists():
        say(f"    ! missing screenshot: {name}")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run()
    try:
        r.add_picture(str(p), width=Inches(6.5))
    except Exception as e:
        say(f"    ! failed to embed {name}: {e}")
        return
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(name.replace("_", " ").replace(".png", ""))
    set_font(cr, "Segoe UI", size=9, italic=True, color=(0x6B, 0x67, 0x5A))

def convert_md_to_docx():
    say(f"[*] converting {GUIDE.name} → {DOCX.name}")
    lines = GUIDE.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # Page setup + default font
    for section in doc.sections:
        section.top_margin    = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin   = Inches(0.8)
        section.right_margin  = Inches(0.8)
    style = doc.styles["Normal"]
    style.font.name = "Segoe UI"
    style.font.size = Pt(10.5)

    # Title page
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run("Saadaa Creative Testing Dashboard")
    set_font(tr, "Segoe UI", size=24, bold=True, color=(0x2F, 0x2C, 0x25))
    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run("User Guidebook")
    set_font(sr, "Segoe UI", size=16, italic=True, color=(0x8A, 0x7B, 0x3A))
    doc.add_paragraph()
    from datetime import date
    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = dp.add_run(f"Generated {date.today().isoformat()}  ·  v2 dashboard")
    set_font(dr, "Segoe UI", size=10.5, color=(0x6B, 0x67, 0x5A))
    doc.add_page_break()

    # Body pass
    i = 0
    current_section = ""
    injected_for_section = set()

    def maybe_inject_images():
        nonlocal current_section
        key = current_section.lower()
        if key in injected_for_section: return
        for skey, imgs in SECTION_IMAGES.items():
            if skey in key:
                for im in imgs:
                    add_screenshot(doc, im)
                break
        injected_for_section.add(key)

    while i < len(lines):
        line = lines[i]

        # Fenced code block
        if line.startswith("```"):
            body = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                body.append(lines[i]); i += 1
            add_shaded_paragraph(doc, "\n".join(body))
            i += 1
            continue

        # Image on its own line
        m = IMG_RE.match(line.strip())
        if m:
            # Ignore — screenshots are injected per section instead
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(line[4:].strip())
            set_font(r, "Segoe UI", size=13, bold=True, color=(0x2F, 0x2C, 0x25))
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(4)
            i += 1; continue
        if line.startswith("## "):
            heading = line[3:].strip()
            p = doc.add_paragraph()
            r = p.add_run(heading)
            set_font(r, "Segoe UI", size=16, bold=True, color=(0x8A, 0x5A, 0x0A))
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after  = Pt(6)
            current_section = heading
            injected_for_section.discard(heading.lower())
            maybe_inject_images()
            i += 1; continue
        if line.startswith("# "):
            i += 1; continue    # title already placed
        if line.startswith("#### "):
            p = doc.add_paragraph()
            r = p.add_run(line[5:].strip())
            set_font(r, "Segoe UI", size=11.5, bold=True, color=(0x4A, 0x45, 0x38))
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)
            i += 1; continue

        # Horizontal rule
        if line.strip() == "---":
            doc.add_paragraph().add_run().add_break()
            i += 1; continue

        # Bulleted list
        if line.startswith("- "):
            add_para_with_inline(doc, line[2:], style="List Bullet")
            i += 1; continue

        # Blank line
        if not line.strip():
            i += 1; continue

        # Plain paragraph (accumulate soft-wrapped lines)
        buf = [line]
        i += 1
        while i < len(lines) and lines[i].strip() and not lines[i].startswith(("- ", "```", "#", "!", "|")):
            buf.append(lines[i])
            i += 1
        add_para_with_inline(doc, " ".join(buf))

    doc.save(str(DOCX))
    say(f"[✓] wrote {DOCX.name}  ({DOCX.stat().st_size // 1024} kB)")

# ────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    build_screenshots()
    convert_md_to_docx()
