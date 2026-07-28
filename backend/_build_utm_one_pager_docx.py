"""Render UTM_Attribution_OnePager.md as a proper Word doc.
Uses python-docx directly (better control over styling than pandoc)."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'UTM_Attribution_OnePager.docx'

doc = Document()

# ── Page margins (tight so it stays 1 page) ──
for s in doc.sections:
    s.top_margin = Cm(1.4); s.bottom_margin = Cm(1.4)
    s.left_margin = Cm(1.6); s.right_margin = Cm(1.6)

# ── Base font ──
style = doc.styles['Normal']
style.font.name = 'Calibri'; style.font.size = Pt(10)

def add_heading(text, size=14, color=(30, 66, 145)):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(size); r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    return p

def add_para(text, size=10, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(3)
    return p

def add_bullet(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs: r.font.size = Pt(10)
    return p

def add_table(headers, rows, first_col_bold=False):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for ri, row in enumerate(rows, start=1):
        cells = t.rows[ri].cells
        for ci, v in enumerate(row):
            cells[ci].text = str(v)
            for p in cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if first_col_bold and ci == 0: r.bold = True
    return t

# ═══════════════ CONTENT ═══════════════
title = doc.add_paragraph()
r = title.add_run('UTM Last-Click Attribution — Matching Logic')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(15, 45, 100)
title.paragraph_format.space_after = Pt(4)

# What we're doing
add_heading('What we\'re doing', 13)
add_para('Every Shopify order carries a "click history" in its URL parameters (UTMs) — the ad the '
         'customer last clicked before buying. We use those UTMs to point each order back to the Meta '
         'ad that earned it. Last click, one order → one ad.')
add_para('Each order carries three URL fields we care about:')
add_table(['Field', 'What it should hold'], [
    ['utm_campaign', 'The campaign name or campaign ID'],
    ['utm_term',     'The adset ID (Meta usually puts a numeric ID here)'],
    ['utm_content',  'The ad — either its numeric ad_id or ad_name'],
], first_col_bold=True)

# Cascade
add_heading('The Cascade — 5 steps, first hit wins', 13)
add_para('Every order is walked through these steps in order. As soon as a step matches, attribution '
         'is locked and we stop.')

steps = [
    ('Step 0 — Manual Overrides (highest priority)',
     'If utm_content contains a phrase we\'ve mapped by hand (e.g. DIVYAYRIAC → this specific active ad), '
     'we route the order to that ad immediately. Used for creative-rename fixes we can\'t recover through '
     'pattern matching. ~13,000 orders currently carry a manual marker.'),
    ('Step 1 — Direct Ad ID Match (most trustworthy)',
     'utm_content is a numeric ad_id (like 120233708339810431) that matches an ad exactly. The Meta URL '
     'template used {{ad.id}} — the gold standard. Unambiguous, one ad, one match.'),
    ('Step 2 — Global Ad Name Match',
     'utm_content isn\'t numeric but matches an ad\'s name globally (exact, fuzzy, or substring). This '
     'happens when the URL template used {{ad.name}}. If exactly one ad has that name → win. If multiple '
     'share it (clones), we prefer the ad with more lifetime spend.'),
    ('Step 3 — Adset Scope + Ad Identifier',
     'utm_term matches a known adset_id (the primary gate for this step). Within that adset, we '
     'identify the specific ad in this order — the first method that resolves wins:\n'
     '  PRIMARY — Asset ID + Adset ID match. utm_content contains a user-managed creative code '
     '(Sep-682, SIF-4442-P1, glam.khush, Oct-740) that is registered for exactly one ad in this '
     'adset. Asset IDs are unique per creative and rename-stable, so this is the most reliable Step 3 '
     'signal and is preferred over every other method below.\n'
     '  FALLBACK 1 — Historical ad_name match. utm_content matches any name this ad ever had inside '
     'the adset (renames handled via ad_name_history).\n'
     '  FALLBACK 2 — Fuzzy / substring / token match. utm_content roughly matches an ad_name in the '
     'adset.\n'
     'If exactly one candidate emerges → attribute. If multiple → split by spend weight.'),
    ('Step 4 — Campaign Scope Only',
     'utm_term doesn\'t match any adset, but utm_campaign matches a campaign. We can identify the campaign '
     'but not the specific ad. Weakest match — used sparingly.'),
    ('Step 5 — Unmatched',
     'Order has UTMs but none resolve to a known Meta ad. Typically: direct traffic, organic search, non-'
     'Meta channels, or Meta URL macros that failed to expand.'),
]
for h, body in steps:
    p = doc.add_paragraph()
    r = p.add_run(h)
    r.bold = True; r.font.size = Pt(10.5); r.font.color.rgb = RGBColor(60, 90, 160)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(1)
    p2 = doc.add_paragraph()
    r = p2.add_run(body); r.font.size = Pt(9.5)
    p2.paragraph_format.left_indent = Cm(0.4); p2.paragraph_format.space_after = Pt(3)

# Sign convention
add_heading('Sign Convention You\'ll See in Reports', 13)
add_table(['Column', 'Meaning'], [
    ['shopify_sales',    '₹ from orders our engine attributed to this ad'],
    ['conv_value',       '₹ Meta\'s own pixel attributed to this ad (usually higher — includes view-through)'],
    ['shop_minus_meta',  'shopify − conv. Positive = Shopify sees more. Negative = Meta over-reports.'],
    ['shop_vs_meta_pct', 'Same as above as a percentage of Meta\'s conv value'],
    ['matched_tier',     'Which of Steps 1–5 fired (or the ad\'s most-common tier)'],
    ['matched_value',    'The exact string that fired the match (ad_id, ad_name, or asset code)'],
], first_col_bold=True)

# Special cases
add_heading('Special Cases We Handle', 13)
special = [
    ('Same-creative clones (2+ ads with identical name and asset_id):',
     'Instead of a coin flip, we split orders across the clones proportional to their lifetime Meta spend. '
     'Higher-spending clone gets the larger share. Marked internally as spend-weight-hashed.'),
    ('Renamed creatives (e.g. MEGHA → DIVYAYRIAC on 22-Apr-2025):',
     'Orders placed after the rename are routed to the new ad; orders before stay with the original. '
     'Historical names are stored per ad so re-attribution runs pick them up automatically.'),
    ('URL macro out of sync with ad name:',
     'Sometimes the URL template emits a short creative slug (FoundersAd_131024) but the ad\'s actual name '
     'is longer (HP+IGP+OFF-RS+IHP+FOU_131024_W). We link the slug → the correct ad via manual override, '
     'and re-attribute the orphaned orders.'),
    ('Missing utm_term (null adset):',
     'If utm_content is a known ad_id → attribute via Step 1 anyway. If it\'s a name-only string with a '
     'unique match → Step 2. Otherwise routed via manual override or left unmatched.'),
]
for h, body in special:
    p = doc.add_paragraph()
    r = p.add_run(h); r.bold = True; r.font.size = Pt(9.5)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(0)
    p2 = doc.add_paragraph()
    r = p2.add_run(body); r.font.size = Pt(9.5)
    p2.paragraph_format.left_indent = Cm(0.4); p2.paragraph_format.space_after = Pt(2)

# Reliable / Weakens
add_heading('What Makes This Reliable', 13)
for b in [
    'Step 1 is unbeatable — direct numeric ID match. Any ad whose URL template uses {{ad.id}} never suffers ambiguity.',
    'Step 3 asset method covers the gap for ads whose URL template uses {{ad.name}} and ends up carrying a slug — as long as the asset code is registered.',
    'Manual markers protect ~13,000 orders across 7 fix categories from being clobbered by re-attribution runs.',
    'The insert semantics of the fetcher mean new order data can only ADD attribution, never overwrite existing decisions.',
]: add_bullet(b)

add_heading('What Weakens It', 13)
for b in [
    'Meta URL templates that use {{ad.name}} instead of {{ad.id}} — every ad rename breaks the last-click link. Currently 127 ads in the last 30 days emit names/slugs instead of IDs.',
    'Multi-clone creatives — when 2+ ads share a name AND asset ID, we can only best-effort split by spend. The true click-owner is unrecoverable without {{ad.id}}.',
    'utm_term missing on the order — kills the adset-scoped Step 3 fast path; we fall back to global name lookup or leave unmatched.',
]: add_bullet(b)

# Footer
f = doc.add_paragraph()
r = f.add_run('One-pager · UTM Last-Click Attribution · Saadaa Creative Testing Dashboard · 2026-07-28')
r.italic = True; r.font.size = Pt(8); r.font.color.rgb = RGBColor(120, 120, 120)
f.paragraph_format.space_before = Pt(12); f.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.save(OUT)
print(f'[ok] wrote {OUT}')
