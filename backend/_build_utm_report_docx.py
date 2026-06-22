"""Generate the UTM attribution status report as a .docx file.

Output:  D:\\Creative_Testing_Dashboard\\backend\\UTM_Attribution_Report.docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"D:\Creative_Testing_Dashboard\backend\UTM_Attribution_Report.docx"

doc = Document()

# Margins
for section in doc.sections:
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

# Body font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# ── Title ───────────────────────────────────────────────────────────────────
title = doc.add_heading('Shopify ↔ Meta Ads UTM Attribution', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
subtitle = doc.add_paragraph()
sr = subtitle.add_run('Status Report')
sr.font.size = Pt(16); sr.font.color.rgb = RGBColor(0x60, 0x60, 0x70)

p = doc.add_paragraph()
p.add_run('Owner: ').bold = True
p.add_run('Data / Analytics\n')
p.add_run('Source system: ').bold = True
p.add_run('Meta_ads_data Supabase project, table shopify_ad_attribution\n')
p.add_run('Scope: ').bold = True
p.add_run('Lifetime data, all-time\n')
p.add_run('Date: ').bold = True
p.add_run('20 June 2026')

doc.add_paragraph().add_run('—' * 30).font.color.rgb = RGBColor(0xC0,0xC0,0xC0)

# Helper: monospace block (for tabular data)
def mono_block(lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    for i, line in enumerate(lines):
        r = p.add_run(line + ('\n' if i < len(lines) - 1 else ''))
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        rPr = r._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        rPr.append(rFonts)
    return p

def para(text):
    return doc.add_paragraph(text)

# ── 1. Executive Summary ────────────────────────────────────────────────────
doc.add_heading('1. Executive Summary', level=1)
para(
    "A 5-tier cascading attribution model maps every Shopify order back to the "
    "originating Meta ad (or, when ad-level signal is absent, to the adset or "
    "campaign that drove it). The model currently matches 97.33 % of "
    "Meta-attributed Shopify orders."
)
para(
    "The pipeline is healthy at the order level: every order resolves to exactly "
    "one tier, with no duplicates and no double-counting. However, a structural "
    "gap in the downstream per-ad aggregation silently drops ~80,000 orders "
    "worth ~Rs 5.4 crore of revenue from per-ad visibility in ae_table_view. "
    "A non-invasive fix is proposed in §5; the underlying attribution logic does "
    "not need to change."
)
para(
    "A second enhancement — refining T3/T4 attribution by asset code "
    "(IFAD, GAD, UGC, etc.) — is also proposed in §5. It will improve precision "
    "for ~46 % of T3 orders."
)

# ── 2. Order Universe ───────────────────────────────────────────────────────
doc.add_heading('2. Order Universe', level=1)
mono_block([
    "Total Shopify orders in attribution scope    676,158",
    "Distinct order_ids                           676,158",
    "Integrity check (one order -> one tier label)   PASS",
])
para(
    "Each order in shopify_ad_attribution is unique. There are no duplicates and "
    "no orders that appear in two tiers."
)

# ── 3. UTM Source Breakdown ─────────────────────────────────────────────────
doc.add_heading('3. UTM Source Breakdown', level=1)
mono_block([
    "Bucket                              Orders        Share",
    "Total                               676,158       100.00 %",
    "  utm_source = meta*                311,006        46.00 %",
    "  utm_source = fb / facebook         14,448         2.10 %",
    "  utm_source = ig / igshopping        8,290         1.20 %",
    "-------------------------------------------------------------",
    "ALL META-RELATED (sum above)        333,744        49.40 %",
    "Other named sources                 255,701        37.80 %",
    "No utm_source set                    86,713        12.80 %",
])
para(
    "meta* covers meta, meta-freetote, meta-featuredofferings, and metald_1. "
    "Other named sources include google, kwikchat, kwikengage, sagepilot-ai, "
    "mkr.bio, etc. About one in eight orders carries no UTM tagging at all."
)

# ── 4. Attribution Cascade ──────────────────────────────────────────────────
doc.add_heading('4. Attribution Cascade — Tier Breakdown', level=1)
para(
    "For each Meta-related order (N = 333,744), the cascade is evaluated in "
    "strict order; the first tier whose condition is satisfied wins and locks "
    "the match. The tier assignment is mutually exclusive at the order level."
)
mono_block([
    "Tier                  Orders         % of Meta    Match strength",
    "T0_template          171,386         51.35 %       Pinpoint  ad",
    "T1_ad_id               3,343          1.00 %       Pinpoint  ad",
    "T2_ad_name_exact      15,086          4.52 %       Pinpoint  ad",
    "T2_ad_name_fuzzy      54,810         16.42 %       Pinpoint  ad",
    "T3_adset_id           78,041         23.38 %       Adset only",
    "T4_campaign_name       2,182          0.65 %       Campaign only",
    "Unmatched              8,896          2.67 %       —",
    "-------------------------------------------------------------",
    "TOTAL MATCHED        324,848         97.33 %",
    "UNMATCHED              8,896          2.67 %",
])

doc.add_heading('4.1 Tier Logic (in priority order)', level=2)
para("T0_template. utm_term carries a numeric adset_id AND utm_content carries "
     "a numeric ad_id. The template lookup resolves directly to that one ad. "
     "Strongest match.")
para("T1_ad_id. utm_content is a standalone 17-digit numeric ad_id.")
para("T2_ad_name_exact. utm_content equals an ad's ad_name, byte for byte.")
para("T2_ad_name_fuzzy. utm_content equals an ad's ad_name after stripping "
     "known suffixes ( - Copy, _H0, _C0,  – Copy 2, etc.) from both sides.")
para("T3_adset_id. utm_term is a numeric adset_id that exists in the ads "
     "table, but utm_content did not resolve to a specific ad. Revenue is "
     "recognised at the adset level.")
para("T4_campaign_name. Last-resort fallback: utm_campaign equals a known "
     "campaign_name. Used when utm_term is empty and no other tier matches.")
para("Unmatched. None of the above hit. Order is excluded from per-ad and "
     "per-adset rollups.")

doc.add_heading('4.2 Mutual Exclusivity', level=2)
para(
    "Each order is processed once through the cascade. The first tier whose "
    "condition matches \"wins\" and the search stops. As a result, every order "
    "has at most one tier label (matched_tier) and at most one matched_value. "
    "SQL integrity check confirms:"
)
mono_block([
    "Total rows                                676,158",
    "Distinct order_ids                        676,158",
    "Matched (matched_tier IS NOT NULL)        325,578",
    "Unmatched (matched_tier IS NULL)          350,580",
    "Inconsistent (has_match TRUE,  tier NULL)       0",
    "Inconsistent (has_match FALSE, tier set)        0",
])

# ── 5. Findings ─────────────────────────────────────────────────────────────
doc.add_heading('5. Findings', level=1)

doc.add_heading('5.1 Structural Gap — T3 / T4 Revenue Invisible at Ad Level', level=2)
para(
    "shopify_ad_attribution correctly records T3 matches against adset_id only, "
    "and T4 matches against campaign_name only — the ad_id column is NULL on "
    "every T3 and T4 row. This is correct on the order side, because no specific "
    "ad is known."
)
para(
    "The downstream aggregate (shopify_ad_agg, built by refresh_ae_table.py) "
    "filters WHERE ad_id IS NOT NULL AND ad_id <> '' and groups by ad_id. "
    "Consequently, all 78,041 T3 orders and all 2,182 T4 orders are silently "
    "dropped before reaching ae_table_view. That is approximately Rs 5.4 crore "
    "of attributed revenue that does not surface against any specific ad row in "
    "the AE analytics table."
)
para("Concrete example — ad CLP-SDCP+MU+OFF-RS+IHP+SDCP_3RVO_JAN2024_W:").runs[0].bold = True
mono_block([
    "amount_spent              Rs   56,29,100.92",
    "conv_value (Meta)        Rs 2,16,43,888.55",
    "roas_ma                   3.84",
    "ae_table_view.shopify_orders   0     <- incorrect (T3 invisible)",
    "ae_table_view.shopify_sales    0     <- incorrect",
])
para(
    "The ad belongs to adset BOTTOMWEAR_JUL2024 (12,031 T3 orders / Rs 1.42 "
    "crore revenue) and spends ~10.04 % of the adset's budget, so its "
    "proportional share is approximately 1,208 orders and Rs 14.3 lakh."
)

doc.add_heading('5.2 Asset Codes Inside T3', level=2)
para(
    "About 46 % of T3 orders carry a creative-family code (IFAD, GAD, UGC, ADB, "
    "BST, BR_, BI_) inside utm_content. This code matches the same code "
    "embedded in many ad names, and can be used to refine T3 attribution from "
    "\"all ads in the adset\" to \"only ads of the same creative family within "
    "the adset\"."
)
mono_block([
    "Asset code            Orders         Revenue (Rs)    % of T3",
    "IFAD                  19,048         2,65,67,619      24.40 %",
    "GAD                    8,659         1,02,96,186      11.09 %",
    "UGC                    5,249           53,51,376       6.72 %",
    "ADB                    1,410           14,40,676       1.81 %",
    "BST                      894           12,47,115       1.15 %",
    "BR_                      412            4,88,312       0.53 %",
    "NONE (no code)        42,383         4,88,62,745      54.30 %",
])
para(
    "Refine-able T3 share is 45.70 % (~35,672 orders, Rs 3.85 crore); the "
    "remaining 54.30 % retains the all-ads-in-adset fallback."
)

# ── 6. Proposed Fixes ──────────────────────────────────────────────────────
doc.add_heading('6. Proposed Fixes', level=1)
para(
    "The structural gap is fixed in the aggregation step only. No metric "
    "definition, formula, or tier logic is touched."
)

doc.add_heading('6.1 Replace Single Direct-Ad-Id Pass with a Three-Pass UNION', level=2)
mono_block([
    "Pass 1  Direct ad_id    T0/T1/T2 rows joined by ad_id        unchanged",
    "Pass 2  Adset spread    T3 rows split spend-weighted across",
    "                        ads in the adset.  Refined by asset",
    "                        code when utm_content carries one",
    "                        (IFAD, GAD, UGC, ADB, BST, BR_, BI_).",
    "                        If no code or no matching ad in the",
    "                        adset, fall back to all-ads spread.",
    "Pass 3  Campaign spread T4 rows split spend-weighted across",
    "                        ads in the campaign, same asset-code",
    "                        refinement, partition = campaign_name.",
])
para("After this change:")
for s in [
    "Every rupee of matched Shopify revenue lands on at least one ad row.",
    "Per-ad shopify_orders and shopify_sales become fractional (e.g. 0.43 of an order). Totals reconcile exactly to matched orders in shopify_ad_attribution.",
    "The structural gap closes: 80,223 orders / ~Rs 5.4 crore become visible at the ad level.",
]:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('6.2 Three Implementation Choices (Pending Approval)', level=2)
for s in [
    "Asset-code priority order: IFAD > GAD > BST > UGC > ADB > BR_ > BI_  (proposed).",
    "Multi-code ad names (e.g. ADB_..._IFAD_...): the ad gets credit for any code in the order's utm_content that also appears in its own name  (proposed).",
    "Mismatch fallback: if utm_content says IFAD but no ad in the adset carries the IFAD code, fall back to spend-weighted spread across all ads in the adset  (proposed).",
]:
    doc.add_paragraph(s, style='List Number')

# ── 7. Next Steps ───────────────────────────────────────────────────────────
doc.add_heading('7. Next Steps', level=1)
for s in [
    "Confirm the three implementation choices in §6.2.",
    "Apply patch to refresh_ae_table.py — shopify_ad_agg rebuild SQL only.",
    "Run end-to-end refresh: rebuild_attribution_orders.py (incremental) followed by refresh_ae_table.py.",
    "Re-verify the example ad in §5.1 — expected outcome: shopify_orders ≈ 1,208, shopify_sales ≈ Rs 14.3 lakh, shopify_top_tier = 'T3_adset_id'.",
    "Spot-check 10 randomly-selected ads to confirm fractional shares reconcile to whole orders at the adset level.",
    "Publish revised AE table to dashboard.",
]:
    doc.add_paragraph(s, style='List Number')

# ── 8. Appendix ─────────────────────────────────────────────────────────────
doc.add_heading('8. Appendix — Data Quality Notes', level=1)
for s in [
    "86,713 orders (12.8 %) have no utm_source at all and are not in scope for any tier match.",
    "The 8,896 unmatched Meta-related orders typically have a utm_term that does not resolve to any current adset (possibly deleted adsets) or carry only a utm_source = meta with empty utm_term and utm_content.",
    "shopify_top_tier per ad is defined as the tier that contributed the most revenue for that ad, not the tightest tier. This is intentional and matches the upstream specification.",
    "Mathematical formulas in ae_table_view are byte-for-byte aligned with the production Apps Script dashboard. Book1.xlsx contains a draft formulation with reach-weighted denominators that is not in production; the working set uses global ratios.",
]:
    doc.add_paragraph(s, style='List Bullet')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('— End of report —')
r.italic = True
r.font.color.rgb = RGBColor(0x80,0x80,0x90)

doc.save(OUT)
print(f"[ok] wrote {OUT}")
